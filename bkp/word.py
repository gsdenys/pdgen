import docx 
import pandas as pd

from docx.shared import Inches, Cm


def __copy_doc(new_name: str):
     with open('model.docx', 'rb') as src, open(new_name, 'wb') as dst:
        dst.write(src.read())
   
def get_file():
    file = "data-dictionary.docx"
    __copy_doc(file)
    
    return docx.Document(file)
   
def addTable(df, doc):     
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    #= 'Colorful Shading Accent 3'

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])     

    t.style = 'Light Grid Accent 1'
    

def to_word(df):
    doc = get_file()
    
    for k, v in df.items():
        doc.add_heading(k.title(), 2)
        addTable(v, doc)

    # save the doc
    doc.save("data-dictionary.docx")